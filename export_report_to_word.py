"""One-off helper: Markdown report -> .docx (requires python-docx)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).resolve().parent
MD = ROOT / "Helmet_Violation_System_Project_Report.md"
OUT = ROOT / "Helmet_Violation_System_Project_Report.docx"


def strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def main():
    if not MD.exists():
        print(f"Missing: {MD}", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### ") and not line.startswith("#### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.lstrip().startswith("|") and "|" in line[1:]:
            buf = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                buf.append(strip_inline_md(lines[i].strip()))
                i += 1
            p = doc.add_paragraph()
            p.add_run("\n".join(buf)).font.name = "Times New Roman"
            p.paragraph_format.space_after = Pt(6)
            continue

        if re.match(r"^[-*] \s*", line):
            doc.add_paragraph(strip_inline_md(re.sub(r"^[-*]\s+", "", line)), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(strip_inline_md(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            i += 1
            continue

        doc.add_paragraph(strip_inline_md(line))
        i += 1

    # Title-ish centering for very first heading if present
    if doc.paragraphs:
        pass

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
